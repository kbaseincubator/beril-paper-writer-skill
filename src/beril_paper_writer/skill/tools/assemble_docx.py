#!/usr/bin/env python3
"""assemble_docx.py — markdown → docx renderer (Tier 2.3 implementation).

Standalone script invoked by `commands/assemble.py`:

    python3 "$SKILL_DIR/tools/assemble_docx.py" <input_md> <output_docx>

Per DECISIONS D-024: python-docx (pure Python) rather than pandoc (system
binary). The package is declared in pyproject.toml runtime deps; no extra
install required. Per the v0.3 punch list, the renderer is hand-rolled
(no markdown library) to keep the dep surface unchanged from v0.2.

Block elements supported:
  - Headings H1-H6 (`# ` … `###### `)
  - Paragraphs (blank-line separated)
  - Bullet lists (`- ` or `* `)
  - Numbered lists (`1. `, `2. `, …)
  - Blockquotes (`> `)
  - Markdown tables (`| col1 | col2 |` with `|---|---|` separator)
  - Fenced code blocks (`` ``` ``)
  - Block-level images on their own line (`![alt](path)`)
  - Horizontal rules (`---`, `***`, `___`)

Inline elements supported (within paragraph / heading / list / cell text):
  - Code spans (`` `text` ``)
  - Bold (`**text**`)
  - Italic (`*text*` or `_text_`)
  - Inline links (`[text](url)`) — rendered as `text (url)` in docx
  - Bare `[N]` citation form preserved verbatim (does NOT match link)

Block-level images are rendered as an inline Picture in a centered
paragraph. Captions come from one of two formats:
  - v0.6.1+: empty alt-text `![]()` followed by a `**Figure N.** ...`
    paragraph — rendered as Caption-styled paragraph (visible in both
    markdown and docx).
  - Legacy: alt-text `![Figure N: caption]()` — rendered as Caption-
    styled paragraph from the alt-text (invisible in markdown rendering).
Inline images within a paragraph are not supported.

Exit codes:
  0  success — output docx written
  1  user error (input markdown not found)
  2  runtime error (python-docx not importable; write failure)
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any


# ---------------------------------------------------------------------------
# Dependency check
# ---------------------------------------------------------------------------


def _check_python_docx() -> int:
    """Return 0 if python-docx is importable; 2 otherwise (with stderr msg)."""
    try:
        import docx  # noqa: F401
    except ImportError:
        sys.stderr.write(
            "Error: python-docx not importable. It is declared in pyproject.toml "
            "runtime deps; reinstall via "
            "`pipx install --force "
            "git+ssh://git@github.com/ArkinLaboratory/beril-paper-writer-skill.git`.\n"
        )
        return 2
    return 0


# ---------------------------------------------------------------------------
# Block parser
# ---------------------------------------------------------------------------


@dataclass
class Block:
    """One parsed markdown block.

    `kind` is one of: h1, h2, h3, h4, h5, h6, paragraph, bullet, number,
    blockquote, table, code_block, image, hr.
    `content` shape varies by kind:
      - heading / paragraph / bullet / number / blockquote: str
      - image: tuple (alt: str, path: str)
      - table / code_block: list[str] of raw lines
      - hr: None
    """

    kind: str
    content: Any


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.+)$")
_NUMBER_RE = re.compile(r"^\d+\.\s+(.+)$")
_BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")
_HR_RE = re.compile(r"^[-*_]{3,}$")
# Block image: a line consisting solely of ![alt](path), nothing else.
# Allow an empty alt and reject paths with characters that would be
# fragile in argparse-y contexts.
_IMAGE_BLOCK_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")
# Markdown table separator: `|---|---|` or `| --- | --- |` or `|:---:|`
_TABLE_SEPARATOR_RE = re.compile(r"^\|[\s:|\-]+\|$")


def parse_blocks(text: str) -> list[Block]:
    """Parse markdown text into a list of Block objects.

    Line-based state machine. Pure function; deterministic; no I/O.
    """
    lines = text.splitlines()
    blocks: list[Block] = []
    i = 0
    para_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal para_lines
        if para_lines:
            blocks.append(Block("paragraph", " ".join(para_lines)))
            para_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Fenced code block (must precede everything else — content is opaque).
        if stripped.startswith("```"):
            flush_paragraph()
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            # i points at the closing fence (or past end if unterminated).
            blocks.append(Block("code_block", code_lines))
            if i < len(lines):
                i += 1  # consume closing fence
            continue

        # Empty line: flush accumulated paragraph.
        if not stripped:
            flush_paragraph()
            i += 1
            continue

        # Heading.
        m = _HEADING_RE.match(line)
        if m:
            flush_paragraph()
            level = min(len(m.group(1)), 6)
            blocks.append(Block(f"h{level}", m.group(2).strip()))
            i += 1
            continue

        # Block-level image (line consists solely of the image tag).
        m = _IMAGE_BLOCK_RE.match(stripped)
        if m:
            flush_paragraph()
            blocks.append(Block("image", (m.group(1), m.group(2))))
            i += 1
            continue

        # Bullet list item.
        m = _BULLET_RE.match(stripped)
        if m:
            flush_paragraph()
            blocks.append(Block("bullet", m.group(1)))
            i += 1
            continue

        # Numbered list item.
        m = _NUMBER_RE.match(stripped)
        if m:
            flush_paragraph()
            blocks.append(Block("number", m.group(1)))
            i += 1
            continue

        # Blockquote.
        m = _BLOCKQUOTE_RE.match(stripped)
        if m:
            flush_paragraph()
            blocks.append(Block("blockquote", m.group(1)))
            i += 1
            continue

        # Markdown table: starts with a `|...|` line, second line is the
        # separator. Without the separator on line+1, treat as paragraph.
        if stripped.startswith("|") and stripped.endswith("|") and i + 1 < len(lines):
            next_stripped = lines[i + 1].strip()
            if _TABLE_SEPARATOR_RE.match(next_stripped):
                flush_paragraph()
                table_lines = [line]
                i += 1
                while i < len(lines):
                    cur = lines[i].strip()
                    if cur.startswith("|") and cur.endswith("|"):
                        table_lines.append(lines[i])
                        i += 1
                    else:
                        break
                blocks.append(Block("table", table_lines))
                continue

        # Horizontal rule.
        if _HR_RE.match(stripped):
            flush_paragraph()
            blocks.append(Block("hr", None))
            i += 1
            continue

        # Default: accumulate as paragraph line.
        para_lines.append(stripped)
        i += 1

    flush_paragraph()
    return blocks


# ---------------------------------------------------------------------------
# Inline parser
# ---------------------------------------------------------------------------

# Order matters: code spans first (consume backtick-delimited regions so
# other tokens don't see chars inside them), then bold (longer), then
# italic (shorter), then inline links. Bare `[N]` citations are NOT
# matched as links (they lack the `(...)` URL part) and pass through as
# plain text.
_INLINE_PATTERN = re.compile(
    r"(?P<code>`[^`\n]+`)"
    r"|(?P<bold>\*\*(?P<bold_text>[^*\n]+)\*\*)"
    r"|(?P<italic>\*(?P<italic_text>[^*\n]+)\*)"
    r"|(?P<italic_us>(?<!\w)_(?P<italic_us_text>[^_\n]+)_(?!\w))"
    r"|(?P<link>\[(?P<link_text>[^\]\n]+)\]\((?P<link_url>[^)\n]+)\))"
)


def tokenize_inline(text: str) -> list[tuple[str, Any]]:
    """Tokenize a string of inline markdown into (kind, payload) tuples.

    Pure function. Used for paragraph / heading / list-item / table-cell
    text. Result feeds `render_inline_runs`.
    """
    tokens: list[tuple[str, Any]] = []
    pos = 0
    for m in _INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            tokens.append(("text", text[pos:m.start()]))
        if m.group("code"):
            tokens.append(("code", m.group("code")[1:-1]))
        elif m.group("bold"):
            tokens.append(("bold", m.group("bold_text")))
        elif m.group("italic"):
            tokens.append(("italic", m.group("italic_text")))
        elif m.group("italic_us"):
            tokens.append(("italic", m.group("italic_us_text")))
        elif m.group("link"):
            tokens.append(("link", (m.group("link_text"), m.group("link_url"))))
        pos = m.end()
    if pos < len(text):
        tokens.append(("text", text[pos:]))
    return tokens


def render_inline_runs(paragraph: Any, text: str) -> None:
    """Render inline-parsed text as runs in a docx paragraph.

    `paragraph` is a python-docx Paragraph object; we add Run objects to
    it via `paragraph.add_run(text)` plus run.bold / run.italic / font
    settings.
    """
    from docx.shared import Pt

    for kind, payload in tokenize_inline(text):
        if kind == "text":
            paragraph.add_run(payload)
        elif kind == "code":
            run = paragraph.add_run(payload)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif kind == "bold":
            run = paragraph.add_run(payload)
            run.bold = True
        elif kind == "italic":
            run = paragraph.add_run(payload)
            run.italic = True
        elif kind == "link":
            link_text, link_url = payload
            # v0.3: render as text + parenthetical URL (no clickable
            # hyperlink). Clickable hyperlink plumbing in python-docx
            # requires manual XML wrangling; defer to v0.4 if needed.
            paragraph.add_run(link_text)
            paragraph.add_run(f" ({link_url})")


# ---------------------------------------------------------------------------
# Block rendering
# ---------------------------------------------------------------------------


def _try_paragraph_with_style(doc: Any, style_name: str, text: str = "") -> Any:
    """Add a paragraph with the named style, falling back to default if absent."""
    try:
        return doc.add_paragraph(text, style=style_name)
    except KeyError:
        return doc.add_paragraph(text)


# v0.4 Phase 3: italic Description paragraph detector. The form is
# `*Description: <text>.*` on a single paragraph (the embed-figures step
# emits this immediately after a Picture). The detector deliberately
# requires the marker prefix to avoid catching arbitrary italic text.
_DESC_PARA_RE = re.compile(r"^\*Description:\s.+\*\s*$", re.DOTALL)

# v0.6.1: visible-caption paragraph — `**Figure N.** Caption text`.
_FIGURE_CAPTION_PARA_RE = re.compile(r"^\*\*Figure\s+\d+\.\*\*\s")


def _is_italic_description_paragraph(content: str) -> bool:
    """True if `content` matches the v0.4 Phase 3 italic Description form.

    Used by render_document to apply Caption style to descriptions that
    follow Pictures. Defensive against multi-line content via re.DOTALL.
    """
    return bool(_DESC_PARA_RE.match(content.strip()))


def _is_figure_caption_paragraph(content: str) -> bool:
    """True if `content` starts with `**Figure N.**` (v0.6.1 format).

    Used by render_document to apply Caption style to the visible-caption
    paragraph that follows an image block in the new format.
    """
    return bool(_FIGURE_CAPTION_PARA_RE.match(content.strip()))


def render_image(doc: Any, alt: str, path: str, base_dir: Path) -> None:
    """Render a block-level image as Picture + Caption-styled paragraph.

    Per Wrinkle B: alt-text is `Figure N: <caption>` form, becomes the
    visible caption text. Picture is centered, sized to ~6 inches wide
    (column-fit for US Letter / A4). Path resolved relative to base_dir.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    # Defensive path validation.
    if path.startswith(("/", "~")):
        sys.stderr.write(
            f"WARN: image path is absolute or home-relative; rejecting: {path}\n"
        )
        para = doc.add_paragraph()
        run = para.add_run(f"[FIGURE PATH REJECTED (absolute): {alt}]")
        run.italic = True
        return
    if ".." in Path(path).parts:
        sys.stderr.write(
            f"WARN: image path contains parent-relative segments; rejecting: {path}\n"
        )
        para = doc.add_paragraph()
        run = para.add_run(f"[FIGURE PATH REJECTED (parent-relative): {alt}]")
        run.italic = True
        return

    full_path = (base_dir / path).resolve()
    if not full_path.is_file():
        # Soft error — emit a placeholder paragraph; do not crash the render.
        sys.stderr.write(f"WARN: image file not found: {full_path}\n")
        para = doc.add_paragraph()
        run = para.add_run(f"[FIGURE MISSING: {alt} at {path}]")
        run.italic = True
        return

    # Picture in its own centered paragraph.
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    try:
        run.add_picture(str(full_path), width=Inches(6))
    except Exception as e:
        # python-docx can raise UnrecognizedImageError for non-image files.
        sys.stderr.write(f"WARN: failed to embed picture {full_path}: {e}\n")
        pic_para.add_run(f"[FIGURE EMBED FAILED: {alt}]").italic = True
        return

    # Caption paragraph below — only when alt-text is non-empty.
    # v0.6.1 format uses empty alt with a separate **Figure N.** paragraph
    # that render_document detects and styles as Caption.
    if alt.strip():
        cap_para = _try_paragraph_with_style(doc, "Caption", alt)
        if cap_para.style.name != "Caption":
            # Fallback styling: italic paragraph.
            for r in cap_para.runs:
                r.italic = True


def render_table(doc: Any, lines: list[str]) -> None:
    """Render a markdown table block as a docx table with grid borders.

    First row is treated as header (bold); separator row dropped; remaining
    rows are data. Cells inline-parsed (bold / italic / code etc.).
    """
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if _TABLE_SEPARATOR_RE.match(stripped):
            continue
        # Split on `|`; drop empty leading/trailing (markdown convention
        # is `| col1 | col2 |` with empty strings before/after).
        cells = [c.strip() for c in stripped.split("|")]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        rows.append(cells)

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        # Some templates don't ship Table Grid; fall back silently.
        pass

    for i, row_cells in enumerate(rows):
        # Pad short rows.
        while len(row_cells) < n_cols:
            row_cells.append("")
        for j, cell_text in enumerate(row_cells):
            cell = table.cell(i, j)
            # python-docx cells start with one empty paragraph; we add to it.
            para = cell.paragraphs[0]
            if i == 0:
                # Header row: bold the entire cell.
                run = para.add_run(cell_text)
                run.bold = True
            else:
                render_inline_runs(para, cell_text)


def render_code_block(doc: Any, lines: list[str]) -> None:
    """Render a fenced code block as a single paragraph with monospace font.

    Soft line breaks (`run.add_break()`) preserve line structure within
    the paragraph. No syntax highlighting.
    """
    from docx.shared import Pt

    para = doc.add_paragraph()
    for idx, line in enumerate(lines):
        run = para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        if idx < len(lines) - 1:
            run.add_break()


# ---------------------------------------------------------------------------
# Top-level renderer
# ---------------------------------------------------------------------------


def render_document(input_md: Path, output_docx: Path) -> int:
    """Render a markdown manuscript to docx via python-docx.

    Block-by-block dispatch over `parse_blocks(text)`. The base directory
    for relative image paths is `input_md.parent`. Output written to
    `output_docx`. Returns 0 on success; raises OSError for I/O failures
    (caller's main() catches and returns 2).
    """
    from docx import Document

    text = input_md.read_text(encoding="utf-8")
    blocks = parse_blocks(text)

    doc = Document()
    base_dir = input_md.parent

    prev_kind: str | None = None
    for block in blocks:
        if block.kind in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(block.kind[1:])
            heading = doc.add_heading("", level=level)
            render_inline_runs(heading, block.content)
        elif block.kind == "paragraph":
            # v0.6.1: `**Figure N.** Caption` paragraph immediately
            # following an image gets Caption style (visible-caption
            # format). v0.4 legacy: italic `*Description: ...*` also
            # gets Caption style for backward compat.
            if prev_kind == "image" and (
                _is_figure_caption_paragraph(block.content)
                or _is_italic_description_paragraph(block.content)
            ):
                para = _try_paragraph_with_style(doc, "Caption")
            else:
                para = doc.add_paragraph()
            render_inline_runs(para, block.content)
        elif block.kind == "image":
            alt, path = block.content
            render_image(doc, alt, path, base_dir)
        elif block.kind == "bullet":
            para = _try_paragraph_with_style(doc, "List Bullet")
            render_inline_runs(para, block.content)
        elif block.kind == "number":
            para = _try_paragraph_with_style(doc, "List Number")
            render_inline_runs(para, block.content)
        elif block.kind == "blockquote":
            para = _try_paragraph_with_style(doc, "Quote")
            if para.style.name != "Quote":
                run = para.add_run(block.content)
                run.italic = True
            else:
                render_inline_runs(para, block.content)
        elif block.kind == "table":
            render_table(doc, block.content)
        elif block.kind == "code_block":
            render_code_block(doc, block.content)
        elif block.kind == "hr":
            # Visual separator: empty paragraph.
            doc.add_paragraph()
        # Unknown block kinds silently dropped (defensive; should not occur).
        prev_kind = block.kind

    doc.save(str(output_docx))
    return 0


# ---------------------------------------------------------------------------
# Argparse harness
# ---------------------------------------------------------------------------


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        prog="assemble_docx.py",
        description=(
            "Render markdown manuscript to docx via python-docx. Hand-rolled "
            "renderer (no markdown library); pure-Python aside from python-docx. "
            "See SKILL_DIR/smoke-test/v0_3_punch_list.md Item 2.3 for the "
            "supported markdown subset."
        ),
    )
    parser.add_argument(
        "input_md",
        help="Path to manuscript.md (concatenated IMRAD markdown).",
    )
    parser.add_argument(
        "output_docx",
        help="Path to write manuscript.docx.",
    )
    args = parser.parse_args(argv)

    rc = _check_python_docx()
    if rc != 0:
        return rc

    input_path = Path(args.input_md).expanduser().resolve()
    output_path = Path(args.output_docx).expanduser().resolve()

    if not input_path.is_file():
        sys.stderr.write(f"Error: input markdown not found: {input_path}\n")
        return 1

    try:
        return render_document(input_path, output_path)
    except OSError as e:
        sys.stderr.write(f"Error writing docx: {e}\n")
        return 2


if __name__ == "__main__":
    sys.exit(main())
