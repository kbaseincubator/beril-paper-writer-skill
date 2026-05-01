"""Tests for tools/assemble_docx.py — markdown→docx renderer (Tier 2.3).

Coverage:
  - Block parsing (headings, paragraphs, lists, tables, code, images,
    blockquotes, hr) on synthetic fixtures.
  - Inline tokenizing (bold, italic, code, links, citations).
  - Image-path defensive handling (absolute/parent-relative rejection;
    missing file → soft placeholder; valid PNG → embedded Picture).
  - Live render against the existing functional_dark_matter draft_1
    manuscript.md is exercised in v0.3 smoke runs, not unit tests
    (manuscript.md may not be present in CI environments).
"""

from __future__ import annotations

import importlib.util
import struct
import sys
import zlib
from pathlib import Path

import pytest


# Locate assemble_docx.py via the package layout (skill/tools/).
_HERE = Path(__file__).resolve()
_REPO_ROOT = _HERE.parents[2]
_TOOLS_DIR = _REPO_ROOT / "src" / "beril_paper_writer" / "skill" / "tools"
_ASSEMBLE_DOCX = _TOOLS_DIR / "assemble_docx.py"


def _load_module():
    """Load assemble_docx as a module for direct function-level testing.

    Per `feedback_importlib_dataclass_gotcha.md`: register in sys.modules
    BEFORE exec_module so @dataclass decorators don't crash with NoneType.
    """
    spec = importlib.util.spec_from_file_location(
        "assemble_docx_under_test", _ASSEMBLE_DOCX
    )
    assert spec is not None
    module = importlib.util.module_from_spec(spec)
    sys.modules["assemble_docx_under_test"] = module
    spec.loader.exec_module(module)  # type: ignore[union-attr]
    return module


@pytest.fixture(scope="module")
def assemble_docx():
    return _load_module()


# ---------------------------------------------------------------------------
# Synthetic-PNG helper (1x1 valid PNG so add_picture works in tests)
# ---------------------------------------------------------------------------


def _write_minimal_png(path: Path) -> None:
    """Write a 1x1 black PNG to `path`. Pure stdlib — no Pillow dep."""
    # Hand-crafted minimal PNG: signature + IHDR + IDAT + IEND.
    sig = b"\x89PNG\r\n\x1a\n"

    def _chunk(tag: bytes, data: bytes) -> bytes:
        crc = zlib.crc32(tag + data) & 0xFFFFFFFF
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", crc)

    # IHDR: 1x1 grayscale, 8-bit, no interlace.
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 0, 0, 0, 0)
    # Raw image data: filter=0 + 1 pixel (single byte 0x00).
    raw = b"\x00\x00"
    idat = zlib.compress(raw)
    iend = b""
    path.write_bytes(sig + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", idat) + _chunk(b"IEND", iend))


# ---------------------------------------------------------------------------
# Block parser tests
# ---------------------------------------------------------------------------


def test_parse_blocks_heading_levels(assemble_docx) -> None:
    text = "# H1\n\n## H2\n\n### H3\n\n#### H4\n"
    blocks = assemble_docx.parse_blocks(text)
    kinds = [b.kind for b in blocks]
    assert kinds == ["h1", "h2", "h3", "h4"]
    assert blocks[0].content == "H1"
    assert blocks[3].content == "H4"


def test_parse_blocks_paragraph_with_blank_separator(assemble_docx) -> None:
    text = "First paragraph.\n\nSecond paragraph.\n"
    blocks = assemble_docx.parse_blocks(text)
    assert len(blocks) == 2
    assert blocks[0].kind == "paragraph"
    assert blocks[0].content == "First paragraph."
    assert blocks[1].content == "Second paragraph."


def test_parse_blocks_paragraph_collapses_consecutive_lines(assemble_docx) -> None:
    """Two non-blank consecutive lines collapse into one paragraph."""
    text = "Line one.\nLine two.\n\nNext para.\n"
    blocks = assemble_docx.parse_blocks(text)
    assert len(blocks) == 2
    assert blocks[0].content == "Line one. Line two."


def test_parse_blocks_bullet_list(assemble_docx) -> None:
    text = "- item one\n- item two\n- item three\n"
    blocks = assemble_docx.parse_blocks(text)
    assert [b.kind for b in blocks] == ["bullet"] * 3
    assert blocks[0].content == "item one"


def test_parse_blocks_numbered_list(assemble_docx) -> None:
    text = "1. first\n2. second\n3. third\n"
    blocks = assemble_docx.parse_blocks(text)
    assert [b.kind for b in blocks] == ["number"] * 3
    assert blocks[2].content == "third"


def test_parse_blocks_image_alone_on_line(assemble_docx) -> None:
    text = "Para one.\n\n![Figure 3: Cross-organism concordance](figures/fig04_concordance.png)\n\nPara two.\n"
    blocks = assemble_docx.parse_blocks(text)
    kinds = [b.kind for b in blocks]
    assert kinds == ["paragraph", "image", "paragraph"]
    alt, path = blocks[1].content
    assert alt == "Figure 3: Cross-organism concordance"
    assert path == "figures/fig04_concordance.png"


def test_parse_blocks_table_with_separator(assemble_docx) -> None:
    text = "| col1 | col2 |\n|---|---|\n| a | b |\n| c | d |\n"
    blocks = assemble_docx.parse_blocks(text)
    assert len(blocks) == 1
    assert blocks[0].kind == "table"
    # Three rows including header (separator dropped at render time).
    assert len(blocks[0].content) == 4


def test_parse_blocks_table_without_separator_is_paragraph(assemble_docx) -> None:
    """A line starting with `|` but no separator on the next line is prose."""
    text = "| this is just a leading pipe.\n"
    blocks = assemble_docx.parse_blocks(text)
    assert blocks[0].kind == "paragraph"


def test_parse_blocks_fenced_code(assemble_docx) -> None:
    text = "Para.\n\n```\nx = 1\ny = 2\n```\n\nNext.\n"
    blocks = assemble_docx.parse_blocks(text)
    kinds = [b.kind for b in blocks]
    assert kinds == ["paragraph", "code_block", "paragraph"]
    assert blocks[1].content == ["x = 1", "y = 2"]


def test_parse_blocks_blockquote(assemble_docx) -> None:
    text = "> This is a quote.\n"
    blocks = assemble_docx.parse_blocks(text)
    assert blocks[0].kind == "blockquote"
    assert blocks[0].content == "This is a quote."


def test_parse_blocks_horizontal_rule(assemble_docx) -> None:
    text = "Above.\n\n---\n\nBelow.\n"
    blocks = assemble_docx.parse_blocks(text)
    kinds = [b.kind for b in blocks]
    assert kinds == ["paragraph", "hr", "paragraph"]


# ---------------------------------------------------------------------------
# Inline tokenizer tests
# ---------------------------------------------------------------------------


def test_tokenize_inline_plain_text(assemble_docx) -> None:
    tokens = assemble_docx.tokenize_inline("just plain text")
    assert tokens == [("text", "just plain text")]


def test_tokenize_inline_bold(assemble_docx) -> None:
    tokens = assemble_docx.tokenize_inline("a **bold** word")
    assert tokens == [("text", "a "), ("bold", "bold"), ("text", " word")]


def test_tokenize_inline_italic(assemble_docx) -> None:
    tokens = assemble_docx.tokenize_inline("an *italic* word")
    assert tokens == [("text", "an "), ("italic", "italic"), ("text", " word")]


def test_tokenize_inline_code_span(assemble_docx) -> None:
    tokens = assemble_docx.tokenize_inline("call `func()` here")
    assert tokens == [("text", "call "), ("code", "func()"), ("text", " here")]


def test_tokenize_inline_link(assemble_docx) -> None:
    tokens = assemble_docx.tokenize_inline("see [docs](https://example.com) for more")
    assert ("link", ("docs", "https://example.com")) in tokens


def test_tokenize_inline_citation_is_text(assemble_docx) -> None:
    """Bare [N] citation form must not match the link regex."""
    tokens = assemble_docx.tokenize_inline("Per Smith [12], this holds.")
    # Whole string is one text token (no link match, no other inline).
    assert tokens == [("text", "Per Smith [12], this holds.")]


def test_tokenize_inline_mixed_bold_italic_code(assemble_docx) -> None:
    text = "a **b** c *d* e `f`"
    tokens = assemble_docx.tokenize_inline(text)
    kinds = [k for k, _ in tokens]
    assert "bold" in kinds
    assert "italic" in kinds
    assert "code" in kinds


# ---------------------------------------------------------------------------
# End-to-end render tests (write actual docx, parse it back)
# ---------------------------------------------------------------------------


def test_render_minimal_document(assemble_docx, tmp_path) -> None:
    """A small heading + paragraph round-trips into a valid docx."""
    md = tmp_path / "in.md"
    md.write_text("# Title\n\nA paragraph with **bold** text.\n", encoding="utf-8")
    out = tmp_path / "out.docx"

    rc = assemble_docx.render_document(md, out)
    assert rc == 0
    assert out.is_file()

    from docx import Document

    doc = Document(str(out))
    # Heading + paragraph = 2 paragraphs.
    assert len(doc.paragraphs) == 2
    assert doc.paragraphs[0].text == "Title"
    assert doc.paragraphs[0].style.name.startswith("Heading")
    # Bold word appears in second paragraph as a bolded run.
    assert "bold" in doc.paragraphs[1].text
    bold_runs = [r for r in doc.paragraphs[1].runs if r.bold]
    assert len(bold_runs) == 1


def test_render_image_block_embeds_picture(assemble_docx, tmp_path) -> None:
    """A block-level image tag becomes a Picture + Caption paragraph."""
    fig_dir = tmp_path / "figures"
    fig_dir.mkdir()
    png_path = fig_dir / "fig01_test.png"
    _write_minimal_png(png_path)

    md = tmp_path / "in.md"
    md.write_text(
        "Some prose (Fig. 1).\n\n"
        "![Figure 1: A test caption](figures/fig01_test.png)\n\n"
        "More prose.\n",
        encoding="utf-8",
    )
    out = tmp_path / "out.docx"
    rc = assemble_docx.render_document(md, out)
    assert rc == 0
    assert out.is_file()

    from docx import Document

    doc = Document(str(out))
    # Inline shapes attached to a Run inside one of the paragraphs.
    n_pictures = len(doc.inline_shapes)
    assert n_pictures == 1, f"expected 1 picture, found {n_pictures}"

    # Caption paragraph contains the alt-text.
    para_texts = [p.text for p in doc.paragraphs]
    assert any("Figure 1: A test caption" in t for t in para_texts)


def test_render_v061_visible_caption_format(assemble_docx, tmp_path) -> None:
    """v0.6.1: empty alt `![](path)` + `**Figure N.** Caption` paragraph
    produces Picture + Caption-styled paragraph in docx."""
    fig_dir = tmp_path / "figures"
    fig_dir.mkdir()
    png_path = fig_dir / "fig01_test.png"
    _write_minimal_png(png_path)

    md = tmp_path / "in.md"
    md.write_text(
        "Some prose (Fig. 1).\n\n"
        "![](figures/fig01_test.png)\n\n"
        "**Figure 1.** A visible caption with detail.\n\n"
        "More prose.\n",
        encoding="utf-8",
    )
    out = tmp_path / "out.docx"
    rc = assemble_docx.render_document(md, out)
    assert rc == 0
    assert out.is_file()

    from docx import Document

    doc = Document(str(out))
    n_pictures = len(doc.inline_shapes)
    assert n_pictures == 1, f"expected 1 picture, found {n_pictures}"

    # Caption paragraph uses Caption style.
    para_texts = [p.text for p in doc.paragraphs]
    para_styles = [p.style.name for p in doc.paragraphs]
    cap_idx = next(
        (i for i, t in enumerate(para_texts) if "Figure 1." in t),
        None,
    )
    assert cap_idx is not None, f"Caption not found: {para_texts}"
    assert para_styles[cap_idx] == "Caption", (
        f"expected Caption style, got {para_styles[cap_idx]!r}"
    )
    assert "A visible caption with detail" in para_texts[cap_idx]

    # No alt-text-based caption (empty alt → no extra caption paragraph).
    assert not any("Figure 1:" in t for t in para_texts)


def test_is_figure_caption_paragraph_matcher(assemble_docx) -> None:
    """v0.6.1: matcher for **Figure N.** visible-caption paragraphs."""
    assert assemble_docx._is_figure_caption_paragraph(
        "**Figure 1.** A caption."
    )
    assert assemble_docx._is_figure_caption_paragraph(
        "**Figure 12.** Multi-sentence caption. Second sentence."
    )
    # Must start with the marker
    assert not assemble_docx._is_figure_caption_paragraph(
        "Figure 1. No bold markers."
    )
    # Tables must not match
    assert not assemble_docx._is_figure_caption_paragraph(
        "**Table 1.** A table caption."
    )


def test_is_italic_description_paragraph_matcher(assemble_docx) -> None:
    """v0.4 Phase 3: matcher used to upgrade post-image italic Description
    paragraphs to Caption style."""
    assert assemble_docx._is_italic_description_paragraph("*Description: Foo.*")
    assert assemble_docx._is_italic_description_paragraph(
        "*Description: Multi-sentence prose. Second sentence.*"
    )
    # Must require the marker prefix
    assert not assemble_docx._is_italic_description_paragraph("*Just italic.*")
    # Must be wrapped in italics, not loose text
    assert not assemble_docx._is_italic_description_paragraph(
        "Description: bare text"
    )


def test_render_description_after_image_gets_caption_style(
    assemble_docx, tmp_path,
) -> None:
    """The italic *Description: ...* paragraph immediately following a
    Picture should be applied with Caption style (visual continuity)."""
    fig_dir = tmp_path / "figures"
    fig_dir.mkdir()
    png_path = fig_dir / "fig01_test.png"
    _write_minimal_png(png_path)

    md = tmp_path / "in.md"
    md.write_text(
        "Some prose (Fig. 1).\n\n"
        "![Figure 1: A caption](figures/fig01_test.png)\n\n"
        "*Description: Growth curves of PA14. Time on x-axis; OD600 on y-axis.*\n\n"
        "More prose.\n",
        encoding="utf-8",
    )
    out = tmp_path / "out.docx"
    rc = assemble_docx.render_document(md, out)
    assert rc == 0

    from docx import Document
    doc = Document(str(out))
    para_styles = [p.style.name for p in doc.paragraphs]
    para_texts = [p.text for p in doc.paragraphs]

    # Find the Description paragraph
    desc_idx = next(
        (i for i, t in enumerate(para_texts) if "Description: Growth curves" in t),
        None,
    )
    assert desc_idx is not None, f"Description not found in paragraphs: {para_texts}"
    assert para_styles[desc_idx] == "Caption", (
        f"expected Caption style on description paragraph, "
        f"got {para_styles[desc_idx]!r}"
    )


def test_description_paragraph_NOT_after_image_stays_normal(
    assemble_docx, tmp_path,
) -> None:
    """An italic Description paragraph that is NOT immediately after an
    image keeps Normal style — the upgrade is image-context-specific."""
    md = tmp_path / "in.md"
    md.write_text(
        "Random italic line not after an image:\n\n"
        "*Description: Free-floating italic text.*\n\n"
        "More prose.\n",
        encoding="utf-8",
    )
    out = tmp_path / "out.docx"
    rc = assemble_docx.render_document(md, out)
    assert rc == 0

    from docx import Document
    doc = Document(str(out))
    para_styles = [p.style.name for p in doc.paragraphs]
    para_texts = [p.text for p in doc.paragraphs]
    desc_idx = next(
        (i for i, t in enumerate(para_texts) if "Free-floating italic text" in t),
        None,
    )
    assert desc_idx is not None
    # Should be Normal (or whatever the default is), NOT Caption
    assert para_styles[desc_idx] != "Caption"


def test_render_image_missing_file_emits_placeholder(assemble_docx, tmp_path) -> None:
    """Missing image → italic placeholder paragraph; render does not crash."""
    md = tmp_path / "in.md"
    md.write_text("![Figure 9: nope](figures/does_not_exist.png)\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    rc = assemble_docx.render_document(md, out)
    assert rc == 0

    from docx import Document

    doc = Document(str(out))
    para_texts = [p.text for p in doc.paragraphs]
    assert any("FIGURE MISSING" in t for t in para_texts)


def test_render_image_absolute_path_rejected(assemble_docx, tmp_path) -> None:
    """Absolute paths → rejection placeholder, not a crash."""
    md = tmp_path / "in.md"
    md.write_text("![Figure X: abs](/etc/passwd)\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    rc = assemble_docx.render_document(md, out)
    assert rc == 0

    from docx import Document

    doc = Document(str(out))
    para_texts = [p.text for p in doc.paragraphs]
    assert any("REJECTED" in t for t in para_texts)


def test_render_table(assemble_docx, tmp_path) -> None:
    md = tmp_path / "in.md"
    md.write_text(
        "| Method | Count |\n|---|---|\n| A | 12 |\n| B | 34 |\n",
        encoding="utf-8",
    )
    out = tmp_path / "out.docx"
    rc = assemble_docx.render_document(md, out)
    assert rc == 0

    from docx import Document

    doc = Document(str(out))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 3  # header + 2 data
    assert table.cell(0, 0).text == "Method"
    assert table.cell(2, 1).text == "34"


def test_render_code_block_uses_monospace(assemble_docx, tmp_path) -> None:
    md = tmp_path / "in.md"
    md.write_text("```\nx = 1\ny = 2\n```\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    rc = assemble_docx.render_document(md, out)
    assert rc == 0

    from docx import Document

    doc = Document(str(out))
    # Code paragraph has runs with monospace font.
    code_paras = [p for p in doc.paragraphs if "x = 1" in p.text]
    assert code_paras
    assert any(r.font.name == "Courier New" for r in code_paras[0].runs)


def test_render_lists(assemble_docx, tmp_path) -> None:
    md = tmp_path / "in.md"
    md.write_text(
        "- bullet one\n- bullet two\n\n1. number one\n2. number two\n",
        encoding="utf-8",
    )
    out = tmp_path / "out.docx"
    rc = assemble_docx.render_document(md, out)
    assert rc == 0

    from docx import Document

    doc = Document(str(out))
    para_texts = [p.text for p in doc.paragraphs]
    assert "bullet one" in para_texts
    assert "number one" in para_texts


# ---------------------------------------------------------------------------
# CLI surface
# ---------------------------------------------------------------------------


def test_main_rejects_missing_input(assemble_docx, tmp_path, capsys) -> None:
    rc = assemble_docx.main(
        [str(tmp_path / "nope.md"), str(tmp_path / "out.docx")]
    )
    assert rc == 1
    captured = capsys.readouterr()
    assert "input markdown not found" in captured.err
